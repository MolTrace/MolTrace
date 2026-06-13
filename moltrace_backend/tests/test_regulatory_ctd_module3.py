"""Prompt 8 — CTD Module 3 section generator.

Drives the real impurity engines (Q3A/B, M7, CPCA) to build the inputs, then cross-checks that the
generated CTD narrative quotes the engine-derived limits (non-circular) and that every number traces
to a SourceRef. Covers exact ICH M4Q section numbering, the DRAFT disclaimer, the three safety-
qualification paths, batch-data traceability, the pharmacopoeial comparison, 3.2.S.3.2, and the
tracked-changes .docx.
"""

from __future__ import annotations

import re

import pytest

from moltrace.regulatory.ctd import (
    CTD_DRAFT_DISCLAIMER,
    CTDSection,
    ImpurityEntry,
    ImpurityOrigin,
    ImpurityProfile,
    SourceKind,
    generate_3p5_impurities,
    generate_3s3_impurities_drug_substance,
)
from moltrace.regulatory.impurities import calculate_q3ab_thresholds, classify_cpca, classify_m7
from moltrace.regulatory.specifications.q6a_builder import BatchResult, SubstanceProfile

_DOSE = 0.5  # g/day (drug product)
_EMS = "CCOS(=O)(=O)C"  # ethyl methanesulfonate — mutagenic (ICH M7 TTC)
_NDMA = "CN(C)N=O"  # N-nitrosodimethylamine — Cohort of Concern (FDA CPCA)


def _profile(**kw) -> ImpurityProfile:
    impurities = (
        ImpurityEntry(
            "Des-methyl examplinib",
            origin=ImpurityOrigin.DEGRADATION,
            observed_levels_percent=(0.05, 0.06, 0.055),
        ),
        ImpurityEntry(
            "Ethyl methanesulfonate",
            structure_smiles=_EMS,
            origin=ImpurityOrigin.PROCESS_RELATED,
            observed_levels_percent=(0.001,),
        ),
        ImpurityEntry(
            "NDMA",
            structure_smiles=_NDMA,
            origin=ImpurityOrigin.PROCESS_RELATED,
            observed_levels_percent=(0.00005,),
        ),
    )
    return ImpurityProfile("Examplinib 250 mg tablets", impurities=impurities, **kw)


def _p55():
    q3ab = calculate_q3ab_thresholds(_DOSE, "drug_product", "oral")
    m7s = [classify_m7(_EMS), classify_m7(_NDMA)]
    cpcas = [classify_cpca(_NDMA)]
    batches = [
        BatchResult("B1", total_impurities_percent=0.30),
        BatchResult("B2", total_impurities_percent=0.35),
        BatchResult("B3", total_impurities_percent=0.28),
    ]
    return generate_3p5_impurities(_profile(), q3ab, m7s, cpcas, batches), q3ab


# --------------------------------------------------------------------------- #
# Structure + ICH M4Q numbering + disclaimer
# --------------------------------------------------------------------------- #
def test_section_numbering_is_exact_m4q() -> None:
    section, _ = _p55()
    assert isinstance(section, CTDSection)
    assert section.section_number == "3.2.P.5.5 / 3.2.P.5.6"
    assert [s.number for s in section.subsections] == ["3.2.P.5.5", "3.2.P.5.6"]
    assert section.m4q_reference == "ICH M4Q(R1)"


def test_draft_disclaimer_is_surfaced() -> None:
    section, _ = _p55()
    assert section.is_draft is True
    assert section.disclaimer == CTD_DRAFT_DISCLAIMER
    md = section.as_markdown()
    assert "DRAFT" in md and "NOT final filing content" in md


# --------------------------------------------------------------------------- #
# 3.2.P.5.5 — the three safety-qualification paths, cross-checked vs the engines
# --------------------------------------------------------------------------- #
def test_ordinary_impurity_uses_q3ab_qualification_threshold() -> None:
    section, q3ab = _p55()
    md = section.as_markdown()
    qual = q3ab.qualification_threshold.effective_percent
    assert f"NMT {qual:g}%" in md  # the engine value, quoted verbatim
    assert "qualification threshold" in md.lower()
    assert any(
        r.kind is SourceKind.Q3AB and "qualification_threshold" in r.locator
        for r in section.sources
    )


def test_mutagenic_impurity_uses_m7_ttc_limit() -> None:
    section, _ = _p55()
    m7 = classify_m7(_EMS)
    expected_ppm = m7.ttc_ug_per_day / _DOSE
    md = section.as_markdown()
    assert f"NMT {expected_ppm:g} ppm" in md
    assert "ICH M7" in md and f"Class {m7.m7_class}" in md
    assert any(r.kind is SourceKind.M7 for r in section.sources)


def test_cohort_of_concern_uses_cpca_ai_limit() -> None:
    section, _ = _p55()
    cpca = classify_cpca(_NDMA)
    expected_ppm = cpca.ai_limit_ng_per_day / _DOSE / 1000.0
    md = section.as_markdown()
    assert f"NMT {expected_ppm:g} ppm" in md
    assert "Cohort-of-Concern" in md and "CPCA" in md
    assert any(r.kind is SourceKind.CPCA for r in section.sources)


def test_m7_path_surfaces_duration_band_and_uses_control_language() -> None:
    # The engine's default-duration TTC is a STAGED (less-than-lifetime) value, not the lifetime
    # 1.5 ug/day; the narrative must surface the duration band and use M7 control (not Q3A/B
    # "qualification") language.
    section, _ = _p55()
    m7 = classify_m7(_EMS)
    md = section.as_markdown()
    assert m7.duration_band in md  # the staged-TTC duration is named, not a bare "TTC"
    assert "staged TTC for the" in md
    assert "Controlled at or below the ICH M7 acceptable intake" in md


def test_dose_denominator_is_traceable() -> None:
    section, _ = _p55()
    assert any(r.kind is SourceKind.Q3AB and r.locator == "daily_dose_g" for r in section.sources)


# --------------------------------------------------------------------------- #
# Traceability — every [S#] marker resolves; batch data is sourced
# --------------------------------------------------------------------------- #
def test_every_marker_resolves_to_a_source() -> None:
    section, _ = _p55()
    index = section.source_index()
    assert index  # non-empty
    # every [S#] used in the narrative paragraphs resolves to an entry
    narrative = " ".join(p for s in section.subsections for p in s.paragraphs)
    used = set(re.findall(r"\[S(\d+)\]", narrative))
    assert used, "expected inline source markers in the narrative"
    for n in used:
        assert f"S{n}" in index
    # all four source kinds are represented
    kinds = {r.kind for r in section.sources}
    assert {SourceKind.Q3AB, SourceKind.M7, SourceKind.CPCA, SourceKind.BATCH} <= kinds


def test_batch_total_impurities_traces_to_batch_data() -> None:
    section, _ = _p55()
    md = section.as_markdown()
    assert "0.35%" in md  # the max total across B1/B2/B3
    assert any(
        r.kind is SourceKind.BATCH and "total_impurities_percent" in r.locator
        for r in section.sources
    )


# --------------------------------------------------------------------------- #
# 3.2.P.5.6 — analytical justification + pharmacopoeial comparison
# --------------------------------------------------------------------------- #
def test_3p56_justifies_method_and_acceptance_criteria() -> None:
    section, _ = _p55()
    p56 = next(s for s in section.subsections if s.number == "3.2.P.5.6")
    assert p56.title == "Justification of Specification(s)"  # exact ICH M4Q heading
    body = " ".join(p56.paragraphs)
    assert "ICH Q2(R2)" in body  # analytical-procedure validation
    assert "stability-indicating" in body
    assert "ICH Q6A" in body  # rationale for acceptance criteria
    assert "total degradation products" in body.lower()  # Q3B scope (not "total impurities")


def test_pharmacopoeial_comparison_when_monograph_supplied() -> None:
    q3ab = calculate_q3ab_thresholds(_DOSE, "drug_product", "oral")
    profile = ImpurityProfile(
        "Examplinib 250 mg tablets",
        impurities=(
            ImpurityEntry(
                "Des-methyl examplinib",
                origin=ImpurityOrigin.DEGRADATION,
                observed_levels_percent=(0.05,),
                pharmacopoeial_limit_percent=0.2,
            ),
        ),
        pharmacopoeial_monograph="Ph. Eur. monograph 9999",
    )
    section = generate_3p5_impurities(profile, q3ab, [], [], [])
    body = " ".join(p for s in section.subsections if s.number == "3.2.P.5.6" for p in s.paragraphs)
    assert "Ph. Eur. monograph 9999" in body
    assert "0.2%" in body and "no less stringent" in body


# --------------------------------------------------------------------------- #
# 3.2.S.3.2 — drug substance
# --------------------------------------------------------------------------- #
def test_3s32_drug_substance_section() -> None:
    q3ab = calculate_q3ab_thresholds(1.0, "drug_substance", "oral")
    substance = SubstanceProfile(
        name="Examplinib", substance_type="drug_substance", max_daily_dose_g=1.0
    )
    profile = ImpurityProfile(
        "Examplinib",
        impurities=(
            ImpurityEntry(
                "Starting material SM-1",
                origin=ImpurityOrigin.PROCESS_RELATED,
                observed_levels_percent=(0.04,),
            ),
            ImpurityEntry(
                "Genotoxic alert imp",
                structure_smiles=_EMS,
                origin=ImpurityOrigin.PROCESS_RELATED,
                observed_levels_percent=(0.0005,),
            ),
        ),
    )
    section = generate_3s3_impurities_drug_substance(substance, profile, q3ab)
    assert section.section_number == "3.2.S.3.2"
    assert section.subsections[0].number == "3.2.S.3.2"
    md = section.as_markdown()
    qual = q3ab.qualification_threshold.effective_percent
    assert f"NMT {qual:g}%" in md
    assert "ICH M7" in md  # structural-alert cross-reference
    assert section.disclaimer == CTD_DRAFT_DISCLAIMER
    assert any(r.kind is SourceKind.Q3AB for r in section.sources)


def test_3s32_pharmacopoeial_limit_is_traceable() -> None:
    q3ab = calculate_q3ab_thresholds(1.0, "drug_substance", "oral")
    substance = SubstanceProfile(
        name="Examplinib", substance_type="drug_substance", max_daily_dose_g=1.0
    )
    profile = ImpurityProfile(
        "Examplinib",
        impurities=(
            ImpurityEntry(
                "Compendial impurity",
                origin=ImpurityOrigin.PROCESS_RELATED,
                observed_levels_percent=(0.08,),
                pharmacopoeial_limit_percent=0.15,
            ),
        ),
        pharmacopoeial_monograph="USP Examplinib",
    )
    section = generate_3s3_impurities_drug_substance(substance, profile, q3ab)
    body = " ".join(section.subsections[0].paragraphs)
    # the compendial limit carries a resolvable [S#] marker (no un-marked threshold)
    assert re.search(r"NMT 0\.15% \[S\d+\]", body)
    assert any(r.kind is SourceKind.PHARMACOPOEIA and "0.15%" in r.value for r in section.sources)


# --------------------------------------------------------------------------- #
# Renderers
# --------------------------------------------------------------------------- #
def test_as_dict_round_trips() -> None:
    section, _ = _p55()
    payload = section.as_dict()
    assert payload["section_number"] == "3.2.P.5.5 / 3.2.P.5.6"
    assert payload["is_draft"] is True
    assert [s["number"] for s in payload["subsections"]] == ["3.2.P.5.5", "3.2.P.5.6"]
    assert payload["sources"] and payload["sources"][0]["id"] == "S1"


def test_to_docx_is_a_tracked_changes_draft(tmp_path) -> None:
    docx = pytest.importorskip("docx")
    from docx.oxml.ns import qn

    section, _ = _p55()
    out = section.to_docx(tmp_path / "ctd_3p55.docx", tracked_changes=True)
    assert out.exists() and out.stat().st_size > 0
    reopened = docx.Document(str(out))
    # tracked changes enabled at the document-settings level
    assert reopened.settings.element.find(qn("w:trackChanges")) is not None
    text = "\n".join(p.text for p in reopened.paragraphs)
    assert "DRAFT" in text and "3.2.P.5.5" in text
