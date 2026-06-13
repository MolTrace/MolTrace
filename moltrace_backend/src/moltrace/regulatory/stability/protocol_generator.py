"""ICH Q1A(R2) stability protocol generator (Prompt 11).

:func:`generate_stability_protocol` turns a product description (dosage form, route, packaging,
intended markets) into a complete, deterministic ICH Q1A(R2) stability protocol:

1. **Climate-zone determination** from the intended markets (WHO/ICH stability zones I–IVb).
2. **Storage conditions** per zone — long-term (zone-specific), intermediate (30 °C/65 % RH, the
   contingency condition for Zone I/II), and accelerated (40 °C/75 % RH).
3. **Testing timepoints** per ICH Q1A(R2): long-term 0/3/6/9/12/18/24/36 months (≥ 12 at
   submission), intermediate 0/6/9/12, accelerated 0/3/6.
4. **Testing matrix** — primary batches × conditions × timepoints, with the bracketing/matrixing
   (ICH Q1D) reduction noted.
5. **Attribute panel** by dosage form (universal attributes plus form-specific ones).
6. **Statistical analysis plan** — ICH Q1E regression for shelf-life estimation.

The protocol renders to Markdown (zero-dependency) via :meth:`StabilityProtocol.as_markdown` or to a
Word ``.docx`` via :meth:`StabilityProtocol.to_docx` (which needs the optional ``python-docx``
dependency: ``pip install 'moltrace[docx]'``). Authoritative constants are encoded from the WHO/ICH
guidelines and cited; this is decision-support — a qualified person reviews and signs the protocol.
"""

from __future__ import annotations

from dataclasses import dataclass
from enum import StrEnum
from pathlib import Path

__all__ = [
    "ClimateZone",
    "ConditionType",
    "StabilityProtocol",
    "StatisticalAnalysisPlan",
    "StorageCondition",
    "TestAttribute",
    "generate_stability_protocol",
]

_GUIDELINE = "ICH Q1A(R2)"
_DISCLAIMER = (
    "Decision-support only — a generated protocol template, NOT a filed stability commitment. The "
    "climate zones, storage conditions, timepoints, and attribute panel are encoded from the "
    "WHO/ICH guidelines and must be reviewed and signed by a qualified "
    "stability/regulatory-affairs professional, and reconciled against the current regional "
    "requirements, before execution or "
    "submission."
)


# --------------------------------------------------------------------------- #
# Climate zones (WHO/ICH stability zones)
# --------------------------------------------------------------------------- #
class ClimateZone(StrEnum):
    """WHO/ICH stability climate zones."""

    ZONE_I = "I"  # temperate
    ZONE_II = "II"  # subtropical / Mediterranean (ICH regions: US, EU, Japan)
    ZONE_III = "III"  # hot, dry
    ZONE_IVA = "IVa"  # hot, humid
    ZONE_IVB = "IVb"  # hot, very humid


# Long-term storage condition (temperature °C, relative humidity %) per zone (WHO/ICH).
# NOTE: Zone III is hot-DRY (30 °C/35 % RH), not 30 °C/65 % RH — the latter is Zone IVa.
_ZONE_LONG_TERM: dict[ClimateZone, tuple[float, float]] = {
    ClimateZone.ZONE_I: (21.0, 45.0),
    ClimateZone.ZONE_II: (25.0, 60.0),
    ClimateZone.ZONE_III: (30.0, 35.0),
    ClimateZone.ZONE_IVA: (30.0, 65.0),
    ClimateZone.ZONE_IVB: (30.0, 75.0),
}

# Stringency rank for ordering / reporting the governing condition.
_ZONE_RANK: dict[ClimateZone, int] = {
    ClimateZone.ZONE_I: 1,
    ClimateZone.ZONE_II: 2,
    ClimateZone.ZONE_III: 3,
    ClimateZone.ZONE_IVA: 4,
    ClimateZone.ZONE_IVB: 5,
}

# Intended market -> climate zone. Keys are upper-cased on lookup.
_MARKET_ZONE: dict[str, ClimateZone] = {
    "US": ClimateZone.ZONE_II,
    "USA": ClimateZone.ZONE_II,
    "EU": ClimateZone.ZONE_II,
    "EUROPE": ClimateZone.ZONE_II,
    "UK": ClimateZone.ZONE_II,
    "JAPAN": ClimateZone.ZONE_II,
    "JP": ClimateZone.ZONE_II,
    "CANADA": ClimateZone.ZONE_II,
    "CHINA": ClimateZone.ZONE_II,
    "KOREA": ClimateZone.ZONE_II,
    "AUSTRALIA": ClimateZone.ZONE_II,
    "GULF": ClimateZone.ZONE_IVA,  # GCC states are Zone IVa (30 °C/65 % RH) per WHO / SFDA
    "MENA": ClimateZone.ZONE_IVA,  # most of MENA is Zone IVa; confirm per country
    "IRAQ": ClimateZone.ZONE_III,  # the documented hot-dry (30 °C/35 % RH) exception (WHO table)
    "ASEAN": ClimateZone.ZONE_IVB,  # ASEAN adopted Zone IVb (30 °C/75 % RH)
    "BRAZIL": ClimateZone.ZONE_IVB,
    "LATAM": ClimateZone.ZONE_IVB,
    "INDIA": ClimateZone.ZONE_IVB,
    "WHO": ClimateZone.ZONE_IVB,
    "INTERNATIONAL": ClimateZone.ZONE_IVB,
}

# ICH Q1A(R2) testing timepoints (months).
_LONG_TERM_TIMEPOINTS = (0, 3, 6, 9, 12, 18, 24, 36)
_INTERMEDIATE_TIMEPOINTS = (0, 6, 9, 12)  # ICH minimum of 4 points over 12 months
_ACCELERATED_TIMEPOINTS = (0, 3, 6)  # ICH minimum of 3 points over 6 months

_ACCELERATED_CONDITION = (40.0, 75.0)
_INTERMEDIATE_CONDITION = (30.0, 65.0)


# --------------------------------------------------------------------------- #
# Dosage-form attribute panels
# --------------------------------------------------------------------------- #
_UNIVERSAL_ATTRIBUTES = (
    "Appearance / description",
    "Assay (content of active)",
    "Degradation products (related substances)",
    "Water content",
)

# Keyword -> (canonical dosage-form label, form-specific attributes).
_DOSAGE_FORM_ATTRIBUTES: dict[str, tuple[str, tuple[str, ...]]] = {
    "tablet": (
        "oral tablet",
        ("Dissolution", "Hardness", "Friability", "Disintegration"),
    ),
    "capsule": (
        "oral capsule",
        (
            "Dissolution",
            "Disintegration",
            "Water activity",
            "Capsule shell integrity / brittleness",
        ),
    ),
    "injectable": (
        "injectable solution",
        (
            "pH",
            "Clarity and colour of solution",
            "Particulate matter (sub-visible)",
            "Sterility",
            "Bacterial endotoxins",
            "Osmolality",
            "Antimicrobial-preservative content (if present)",
            "Container-closure integrity",
        ),
    ),
    "solution": (
        "solution",
        ("pH", "Clarity and colour of solution", "Osmolality", "Particulate matter (sub-visible)"),
    ),
    "nasal": (
        "nasal spray",
        (
            "pH",
            "Droplet / particle size distribution",
            "Spray pattern",
            "Plume geometry",
            "Delivered-dose (shot-weight) uniformity",
            "Pump delivery / actuation force",
            "Microbial limits",
            "Antimicrobial-preservative content (if present)",
            "Leachables",
        ),
    ),
    "transdermal": (
        "transdermal patch",
        (
            "In-vitro drug-release rate",
            "Adhesion / peel / tack",
            "Cold flow",
            "Crystallization (microscopy)",
            "Microbial limits",
            "Leachables",
        ),
    ),
    "patch": (
        "transdermal patch",
        (
            "In-vitro drug-release rate",
            "Adhesion / peel / tack",
            "Cold flow",
            "Crystallization (microscopy)",
            "Microbial limits",
            "Leachables",
        ),
    ),
}

# Routes that imply a sterile product (add sterility + endotoxins if not already present).
_STERILE_ROUTES = frozenset(
    {
        "parenteral",
        "injection",
        "injectable",
        "intravenous",
        "iv",
        "im",
        "sc",
        "subcutaneous",
        "intramuscular",
        "ophthalmic",
        "intrathecal",
    }
)


# --------------------------------------------------------------------------- #
# Dataclasses
# --------------------------------------------------------------------------- #
class ConditionType(StrEnum):
    LONG_TERM = "long_term"
    INTERMEDIATE = "intermediate"
    ACCELERATED = "accelerated"


@dataclass(frozen=True)
class StorageCondition:
    """One stability storage condition with its testing timepoints."""

    condition_type: ConditionType
    temperature_c: float
    relative_humidity_percent: float
    timepoints_months: tuple[int, ...]
    zone: ClimateZone | None = None
    tolerance: str = "± 2 °C / ± 5 % RH"
    note: str = ""

    def label(self) -> str:
        return f"{self.temperature_c:g} °C / {self.relative_humidity_percent:g} % RH"


@dataclass(frozen=True)
class TestAttribute:
    """A stability-indicating attribute to test."""

    name: str
    category: str  # "universal" | "dosage-form-specific"


@dataclass(frozen=True)
class StatisticalAnalysisPlan:
    """ICH Q1E statistical analysis plan for shelf-life estimation."""

    method: str
    shelf_life_rule: str
    poolability_test: str
    references: tuple[str, ...]


@dataclass(frozen=True)
class StabilityProtocol:
    """A complete ICH Q1A(R2) stability protocol, ready for QA review.

    ``as_markdown()`` renders the protocol document with zero extra dependencies; ``to_docx(path)``
    renders a Word document (requires the optional ``python-docx`` dependency). ``as_dict()`` is the
    structured payload for persistence or an API.
    """

    dosage_form: str
    route: str
    drug_product_type: str
    package_primary: str
    package_secondary: str
    intended_markets: tuple[str, ...]
    climate_zones: tuple[ClimateZone, ...]
    storage_conditions: tuple[StorageCondition, ...]
    attributes: tuple[TestAttribute, ...]
    n_primary_batches: int
    statistical_plan: StatisticalAnalysisPlan
    guideline: str = _GUIDELINE
    notes: tuple[str, ...] = ()
    warnings: tuple[str, ...] = ()
    disclaimer: str = _DISCLAIMER
    human_review_required: bool = True

    @property
    def long_term_conditions(self) -> tuple[StorageCondition, ...]:
        return tuple(
            c for c in self.storage_conditions if c.condition_type is ConditionType.LONG_TERM
        )

    def as_dict(self) -> dict:
        return {
            "guideline": self.guideline,
            "dosage_form": self.dosage_form,
            "route": self.route,
            "drug_product_type": self.drug_product_type,
            "packaging": {"primary": self.package_primary, "secondary": self.package_secondary},
            "intended_markets": list(self.intended_markets),
            "climate_zones": [z.value for z in self.climate_zones],
            "storage_conditions": [
                {
                    "type": c.condition_type.value,
                    "zone": c.zone.value if c.zone else None,
                    "temperature_c": c.temperature_c,
                    "relative_humidity_percent": c.relative_humidity_percent,
                    "tolerance": c.tolerance,
                    "timepoints_months": list(c.timepoints_months),
                    "note": c.note,
                }
                for c in self.storage_conditions
            ],
            "attributes": [{"name": a.name, "category": a.category} for a in self.attributes],
            "n_primary_batches": self.n_primary_batches,
            "statistical_plan": {
                "method": self.statistical_plan.method,
                "shelf_life_rule": self.statistical_plan.shelf_life_rule,
                "poolability_test": self.statistical_plan.poolability_test,
                "references": list(self.statistical_plan.references),
            },
            "notes": list(self.notes),
            "warnings": list(self.warnings),
            "disclaimer": self.disclaimer,
            "human_review_required": self.human_review_required,
        }

    def as_markdown(self) -> str:
        lines: list[str] = [
            f"# Stability Protocol — {self.guideline}",
            "",
            f"- **Drug-product type:** {self.drug_product_type}",
            f"- **Dosage form:** {self.dosage_form}",
            f"- **Route of administration:** {self.route}",
            f"- **Primary packaging:** {self.package_primary}",
            f"- **Secondary packaging:** {self.package_secondary}",
            f"- **Intended markets:** {', '.join(self.intended_markets)}",
            f"- **Climate zone(s):** {', '.join(z.value for z in self.climate_zones)}",
            f"- **Primary batches:** {self.n_primary_batches} (ICH Q1A(R2) minimum)",
            "- **Status:** pending QA review",
            "",
            "## 1. Storage conditions & testing timepoints",
            "",
            "| Condition | Zone | Storage | Tolerance | Timepoints (months) |",
            "|---|---|---|---|---|",
        ]
        for c in self.storage_conditions:
            tps = ", ".join(str(t) for t in c.timepoints_months)
            zone = c.zone.value if c.zone else "—"
            lines.append(
                f"| {c.condition_type.value} | {zone} | {c.label()} | {c.tolerance} | {tps} |"
            )
        for c in self.storage_conditions:
            if c.note:
                lines.append(f"\n> {c.condition_type.value}: {c.note}")
        lines += ["", "## 2. Attributes to test", ""]
        lines += [f"- {a.name} ({a.category})" for a in self.attributes]
        lines += [
            "",
            "## 3. Testing matrix",
            "",
            f"Each of the {self.n_primary_batches} primary batches is tested for every attribute "
            "at every timepoint of every storage condition above (full design). A reduced design "
            "(bracketing / matrixing, ICH Q1D) may be justified for multiple strengths or "
            "container sizes.",
            "",
            "## 4. Statistical analysis plan (ICH Q1E)",
            "",
            f"- **Method:** {self.statistical_plan.method}",
            f"- **Shelf-life rule:** {self.statistical_plan.shelf_life_rule}",
            f"- **Batch poolability:** {self.statistical_plan.poolability_test}",
            f"- **References:** {', '.join(self.statistical_plan.references)}",
            "",
            "## 5. Notes",
            "",
        ]
        lines += [f"- {n}" for n in self.notes]
        if self.warnings:
            lines += ["", "## 6. Warnings", ""]
            lines += [f"- {w}" for w in self.warnings]
        lines += [
            "",
            "## 7. QA review & approval",
            "",
            f"_{self.disclaimer}_",
            "",
            "- **Prepared by:** ____________________  **Date:** __________",
            "- **Reviewed by (QA):** ____________________  **Date:** __________",
            "- **Approved by:** ____________________  **Date:** __________",
        ]
        return "\n".join(lines)

    def to_docx(self, path: str | Path) -> Path:
        """Render the protocol to a Word ``.docx`` document at *path* and return the path.

        Requires the optional ``python-docx`` dependency (``pip install 'moltrace[docx]'``).
        """

        try:
            from docx import Document
            from docx.shared import Pt
        except ModuleNotFoundError as exc:  # pragma: no cover - exercised via the extra
            raise ModuleNotFoundError(
                "to_docx() requires python-docx; install the optional extra: "
                "pip install 'moltrace[docx]'"
            ) from exc

        doc = Document()
        doc.add_heading(f"Stability Protocol — {self.guideline}", level=0)

        meta = doc.add_table(rows=0, cols=2)
        meta.style = "Light Grid Accent 1"
        for key, val in (
            ("Drug-product type", self.drug_product_type),
            ("Dosage form", self.dosage_form),
            ("Route of administration", self.route),
            ("Primary packaging", self.package_primary),
            ("Secondary packaging", self.package_secondary),
            ("Intended markets", ", ".join(self.intended_markets)),
            ("Climate zone(s)", ", ".join(z.value for z in self.climate_zones)),
            ("Primary batches", f"{self.n_primary_batches} (ICH Q1A(R2) minimum)"),
            ("Status", "pending QA review"),
        ):
            row = meta.add_row().cells
            row[0].text = key
            row[1].text = str(val)

        doc.add_heading("1. Storage conditions & testing timepoints", level=1)
        cond = doc.add_table(rows=1, cols=5)
        cond.style = "Light Grid Accent 1"
        hdr = cond.rows[0].cells
        for i, h in enumerate(("Condition", "Zone", "Storage", "Tolerance", "Timepoints (months)")):
            hdr[i].text = h
        for c in self.storage_conditions:
            cells = cond.add_row().cells
            cells[0].text = c.condition_type.value
            cells[1].text = c.zone.value if c.zone else "—"
            cells[2].text = c.label()
            cells[3].text = c.tolerance
            cells[4].text = ", ".join(str(t) for t in c.timepoints_months)
        for c in self.storage_conditions:
            if c.note:
                doc.add_paragraph(f"{c.condition_type.value}: {c.note}", style="Intense Quote")

        doc.add_heading("2. Attributes to test", level=1)
        for a in self.attributes:
            doc.add_paragraph(f"{a.name} ({a.category})", style="List Bullet")

        doc.add_heading("3. Testing matrix", level=1)
        doc.add_paragraph(
            f"Each of the {self.n_primary_batches} primary batches is tested for every attribute "
            "at every timepoint of every storage condition above (full design). A reduced design "
            "(bracketing / matrixing, ICH Q1D) may be justified for multiple strengths or "
            "container sizes."
        )

        doc.add_heading("4. Statistical analysis plan (ICH Q1E)", level=1)
        for label, val in (
            ("Method", self.statistical_plan.method),
            ("Shelf-life rule", self.statistical_plan.shelf_life_rule),
            ("Batch poolability", self.statistical_plan.poolability_test),
            ("References", ", ".join(self.statistical_plan.references)),
        ):
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: ")
            run.bold = True
            p.add_run(val)

        doc.add_heading("5. Notes", level=1)
        for n in self.notes:
            doc.add_paragraph(n, style="List Bullet")
        if self.warnings:
            doc.add_heading("6. Warnings", level=1)
            for w in self.warnings:
                doc.add_paragraph(w, style="List Bullet")

        doc.add_heading("7. QA review & approval", level=1)
        disc = doc.add_paragraph()
        disc.add_run(self.disclaimer).italic = True
        for role in ("Prepared by", "Reviewed by (QA)", "Approved by"):
            doc.add_paragraph(f"{role}: ____________________    Date: __________")

        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.size is None:
                    run.font.size = Pt(11)

        out = Path(path)
        doc.save(out)
        return out


# --------------------------------------------------------------------------- #
# Generation
# --------------------------------------------------------------------------- #
def _zones_for_markets(markets: list[str]) -> tuple[list[ClimateZone], list[str]]:
    """Map markets to distinct climate zones; unknown markets default to Zone IVb (worst case)."""

    zones: list[ClimateZone] = []
    warnings: list[str] = []
    for market in markets:
        zone = _MARKET_ZONE.get(market.strip().upper())
        if zone is None:
            zone = ClimateZone.ZONE_IVB
            warnings.append(
                f"market '{market}' not in the zone map — assumed Zone IVb (30 °C/75 % RH, worst "
                "case); confirm the correct climate zone"
            )
        if zone not in zones:
            zones.append(zone)
    zones.sort(key=lambda z: _ZONE_RANK[z])
    return zones, warnings


def _attributes_for(dosage_form: str, route: str) -> tuple[list[TestAttribute], list[str]]:
    form_key = dosage_form.strip().lower()
    matched: tuple[str, tuple[str, ...]] | None = None
    for keyword, spec in _DOSAGE_FORM_ATTRIBUTES.items():
        if keyword in form_key:
            matched = spec
            break

    warnings: list[str] = []
    specific: list[str] = []
    if matched is None:
        warnings.append(
            f"dosage form '{dosage_form}' not in the attribute map — only universal attributes "
            "enumerated; add form-specific attributes on review"
        )
    else:
        specific = list(matched[1])

    # Sterile routes require sterility + endotoxins even if the form map didn't add them.
    if route.strip().lower() in _STERILE_ROUTES:
        for attr in ("Sterility", "Bacterial endotoxins"):
            if attr not in specific:
                specific.append(attr)

    attributes = [TestAttribute(name, "universal") for name in _UNIVERSAL_ATTRIBUTES]
    attributes += [TestAttribute(name, "dosage-form-specific") for name in specific]
    return attributes, warnings


def _statistical_plan() -> StatisticalAnalysisPlan:
    return StatisticalAnalysisPlan(
        method=(
            "Linear regression of each quantitative, stability-indicating attribute (assay, "
            "individual and total degradation products, dissolution where quantitative) against "
            "time at the long-term condition; data transformation (e.g. logarithmic) applied only "
            "where justified."
        ),
        shelf_life_rule=(
            "The proposed shelf life is the earliest time at which the 95 % one-sided confidence "
            "limit for the mean attribute response intersects the acceptance criterion (two-sided "
            "95 % for attributes that can change in either direction)."
        ),
        poolability_test=(
            "Batches are tested for poolability by analysis of covariance at the 0.25 significance "
            "level (slopes then intercepts); poolable batches are combined, otherwise the shelf "
            "life is set by the worst-case individual batch."
        ),
        references=("ICH Q1E", "ICH Q1A(R2) §2.2.9 (Evaluation)"),
    )


def _build_notes(zones: list[ClimateZone], has_intermediate: bool) -> list[str]:
    notes = [
        "At least 12 months of long-term data are required at submission; long-term testing "
        "continues through the proposed shelf life (every 3 months in year 1, every 6 months in "
        "year 2, annually thereafter).",
        "Accelerated condition is 40 °C/75 % RH; a 'significant change' (ICH Q1A(R2) §2.2.7.1) at "
        "the accelerated condition triggers testing at the intermediate condition.",
        "Photostability testing on one batch per ICH Q1B.",
        "Bracketing and/or matrixing designs per ICH Q1D may reduce the number of samples for "
        "multiple strengths, container sizes, or fills, with justification.",
        "A minimum of three primary batches is studied; for a drug substance/product manufactured "
        "at production scale, batch selection follows ICH Q1A(R2) §2.1.3.",
    ]
    if has_intermediate:
        notes.append(
            "Intermediate condition (30 °C/65 % RH) is included because a Zone I/II long-term "
            "condition (25 °C/60 % RH) is in scope; it is tested only if a significant change "
            "occurs at the accelerated condition."
        )
    if ClimateZone.ZONE_IVB in zones:
        notes.append("Zone IVb markets (e.g. ASEAN) require the 30 °C/75 % RH long-term condition.")
    return notes


def generate_stability_protocol(
    dosage_form: str,
    route: str,
    package_primary: str,
    package_secondary: str,
    intended_markets: list[str],
    drug_product_type: str = "finished",
) -> StabilityProtocol:
    """Generate an ICH Q1A(R2) stability protocol for a product (see module docstring).

    The protocol covers a long-term condition for each distinct climate zone of the intended
    markets, the accelerated condition (40 °C/75 % RH), and — when a Zone I/II market is in scope —
    the intermediate condition (30 °C/65 % RH) as the significant-change contingency.
    """

    if not intended_markets:
        raise ValueError("at least one intended market is required")

    zones, zone_warnings = _zones_for_markets(list(intended_markets))

    conditions: list[StorageCondition] = []
    for zone in zones:
        temp, rh = _ZONE_LONG_TERM[zone]
        conditions.append(
            StorageCondition(
                condition_type=ConditionType.LONG_TERM,
                temperature_c=temp,
                relative_humidity_percent=rh,
                timepoints_months=_LONG_TERM_TIMEPOINTS,
                zone=zone,
            )
        )

    has_intermediate = any(z in (ClimateZone.ZONE_I, ClimateZone.ZONE_II) for z in zones)
    if has_intermediate:
        temp, rh = _INTERMEDIATE_CONDITION
        conditions.append(
            StorageCondition(
                condition_type=ConditionType.INTERMEDIATE,
                temperature_c=temp,
                relative_humidity_percent=rh,
                timepoints_months=_INTERMEDIATE_TIMEPOINTS,
                note=(
                    "tested only if a significant change occurs at the accelerated condition "
                    "(ICH Q1A(R2) §2.2.7.1)"
                ),
            )
        )

    acc_temp, acc_rh = _ACCELERATED_CONDITION
    conditions.append(
        StorageCondition(
            condition_type=ConditionType.ACCELERATED,
            temperature_c=acc_temp,
            relative_humidity_percent=acc_rh,
            timepoints_months=_ACCELERATED_TIMEPOINTS,
        )
    )

    attributes, attr_warnings = _attributes_for(dosage_form, route)
    notes = _build_notes(zones, has_intermediate)

    return StabilityProtocol(
        dosage_form=dosage_form,
        route=route,
        drug_product_type=drug_product_type,
        package_primary=package_primary,
        package_secondary=package_secondary,
        intended_markets=tuple(intended_markets),
        climate_zones=tuple(zones),
        storage_conditions=tuple(conditions),
        attributes=tuple(attributes),
        n_primary_batches=3,
        statistical_plan=_statistical_plan(),
        notes=tuple(notes),
        warnings=tuple(zone_warnings + attr_warnings),
    )
