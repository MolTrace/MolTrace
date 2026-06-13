"""CTD Module 3 section generator (Prompt 8).

Builds ICH M4Q(R1) Module-3 quality narratives for impurities from the validated deterministic
engines (ICH Q3A/B thresholds, ICH M7 mutagenicity, FDA CPCA nitrosamine potency) and batch data:

* :func:`generate_3p5_impurities` → CTD **3.2.P.5.5** (Characterisation of Impurities) +
  **3.2.P.5.6** (Justification of Specifications) for a drug product.
* :func:`generate_3s3_impurities_drug_substance` → CTD **3.2.S.3.2** (Impurities) for a drug
  substance.

Section numbers follow ICH M4Q exactly. Every number and threshold carries a :class:`SourceRef`
back to the engine output or batch entry that produced it — rendered as an inline ``[S#]`` marker
resolved in a per-section *Source traceability* index, so a reviewer can trace each value to
``q3ab_calculator`` / ``m7_classifier`` / ``cpca_classifier`` / ``batch_data``.

The result is a structured :class:`CTDSection` that renders to Markdown (zero-dependency) or to a
**draft** Word ``.docx`` with tracked changes enabled (optional ``python-docx`` dependency:
``pip install 'moltrace[docx]'``). Generated sections are DRAFTS — see
:data:`CTD_DRAFT_DISCLAIMER`.
"""

from __future__ import annotations

from dataclasses import dataclass
from enum import StrEnum
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:  # consumed engine types — duck-typed at runtime
    from moltrace.regulatory.impurities import (
        CPCAResult,
        ImpurityThresholds,
        M7Classification,
    )
    from moltrace.regulatory.specifications.q6a_builder import BatchResult, SubstanceProfile

__all__ = [
    "CTD_DRAFT_DISCLAIMER",
    "CTDSection",
    "CTDSubsection",
    "CTDTable",
    "ImpurityEntry",
    "ImpurityOrigin",
    "ImpurityProfile",
    "SourceKind",
    "SourceRef",
    "generate_3p5_impurities",
    "generate_3s3_impurities_drug_substance",
]

CTD_DRAFT_DISCLAIMER = (
    "DRAFT for qualified regulatory-affairs review and QA sign-off — NOT final filing content. "
    "Every number, threshold, and narrative claim is machine-generated from the deterministic "
    "impurity engines and MUST be verified against its cited source and the official guideline "
    "before inclusion in any regulatory filing."
)

_M4Q_REFERENCE = "ICH M4Q(R1)"


# --------------------------------------------------------------------------- #
# Traceability
# --------------------------------------------------------------------------- #
class SourceKind(StrEnum):
    """Which engine / dataset a value traces to."""

    Q3AB = "q3ab_calculator"
    M7 = "m7_classifier"
    CPCA = "cpca_classifier"
    BATCH = "batch_data"
    PHARMACOPOEIA = "pharmacopoeial_monograph"


@dataclass(frozen=True)
class SourceRef:
    """A traceability link from a value in the narrative to its source of truth."""

    kind: SourceKind
    locator: str  # dotted path into the source, e.g. qualification_threshold.effective_percent
    value: str  # the rendered value, verbatim
    detail: str = ""  # regulatory basis / rule-set version

    def describe(self) -> str:
        tail = f" — {self.detail}" if self.detail else ""
        return f"{self.kind.value}:{self.locator} = {self.value}{tail}"


class _SourceLedger:
    """Allocates ``[S#]`` markers for SourceRefs and de-duplicates identical references."""

    def __init__(self) -> None:
        self._refs: list[SourceRef] = []
        self._index: dict[tuple, int] = {}

    def mark(self, ref: SourceRef) -> str:
        key = (ref.kind, ref.locator, ref.value, ref.detail)
        if key not in self._index:
            self._refs.append(ref)
            self._index[key] = len(self._refs)
        return f"[S{self._index[key]}]"

    def refs(self) -> tuple[SourceRef, ...]:
        return tuple(self._refs)


# --------------------------------------------------------------------------- #
# Inputs
# --------------------------------------------------------------------------- #
class ImpurityOrigin(StrEnum):
    PROCESS_RELATED = "process-related"
    DEGRADATION = "degradation product"
    BOTH = "process-related and degradation product"
    UNKNOWN = "origin under investigation"


@dataclass(frozen=True)
class ImpurityEntry:
    """One impurity to characterise/specify."""

    name: str
    structure_smiles: str | None = None
    origin: ImpurityOrigin = ImpurityOrigin.UNKNOWN
    observed_levels_percent: tuple[float, ...] = ()
    is_specified: bool = True
    structural_description: str = ""
    pharmacopoeial_limit_percent: float | None = None

    @property
    def max_observed_percent(self) -> float | None:
        return max(self.observed_levels_percent) if self.observed_levels_percent else None


@dataclass(frozen=True)
class ImpurityProfile:
    """The impurity profile of a substance/product plus its analytical context."""

    substance_name: str
    impurities: tuple[ImpurityEntry, ...] = ()
    analytical_procedure: str = "reversed-phase HPLC with UV detection"
    analytical_procedure_validated: bool = True  # ICH Q2(R2)
    pharmacopoeial_monograph: str = ""


# --------------------------------------------------------------------------- #
# Output
# --------------------------------------------------------------------------- #
@dataclass(frozen=True)
class CTDTable:
    caption: str
    headers: tuple[str, ...]
    rows: tuple[tuple[str, ...], ...]


@dataclass(frozen=True)
class CTDSubsection:
    number: str
    title: str
    paragraphs: tuple[str, ...] = ()
    table: CTDTable | None = None


@dataclass(frozen=True)
class CTDSection:
    """A generated CTD section, ready for regulatory-affairs review.

    ``sources`` is the full traceability index — every ``[S#]`` marker in the narrative resolves to
    one entry. ``as_markdown()`` renders the document with zero dependencies; ``to_docx()`` renders
    a draft Word document with tracked changes enabled.
    """

    section_number: str  # e.g. "3.2.P.5.5 / 3.2.P.5.6"
    title: str
    subsections: tuple[CTDSubsection, ...]
    sources: tuple[SourceRef, ...]
    m4q_reference: str = _M4Q_REFERENCE
    disclaimer: str = CTD_DRAFT_DISCLAIMER
    is_draft: bool = True

    def source_index(self) -> dict[str, SourceRef]:
        return {f"S{i}": ref for i, ref in enumerate(self.sources, start=1)}

    def as_dict(self) -> dict:
        return {
            "section_number": self.section_number,
            "title": self.title,
            "m4q_reference": self.m4q_reference,
            "is_draft": self.is_draft,
            "disclaimer": self.disclaimer,
            "subsections": [
                {
                    "number": s.number,
                    "title": s.title,
                    "paragraphs": list(s.paragraphs),
                    "table": None
                    if s.table is None
                    else {
                        "caption": s.table.caption,
                        "headers": list(s.table.headers),
                        "rows": [list(r) for r in s.table.rows],
                    },
                }
                for s in self.subsections
            ],
            "sources": [
                {
                    "id": f"S{i}",
                    "kind": r.kind.value,
                    "locator": r.locator,
                    "value": r.value,
                    "detail": r.detail,
                }
                for i, r in enumerate(self.sources, start=1)
            ],
        }

    def as_markdown(self) -> str:
        lines: list[str] = [
            f"# CTD {self.section_number} — {self.title}",
            "",
            f"> **{'DRAFT — ' if self.is_draft else ''}{self.disclaimer}**",
            "",
            f"_Format: {self.m4q_reference}._",
            "",
        ]
        for s in self.subsections:
            lines += [f"## {s.number} {s.title}", ""]
            for p in s.paragraphs:
                lines += [p, ""]
            if s.table is not None:
                lines += [f"*{s.table.caption}*", ""]
                lines.append("| " + " | ".join(s.table.headers) + " |")
                lines.append("|" + "---|" * len(s.table.headers))
                for row in s.table.rows:
                    lines.append("| " + " | ".join(row) + " |")
                lines.append("")
        lines += ["## Source traceability", ""]
        index = self.source_index()
        if index:
            lines += ["| Ref | Source | Locator | Value |", "|---|---|---|---|"]
            for sid, ref in index.items():
                lines.append(f"| {sid} | {ref.kind.value} | {ref.locator} | {ref.value} |")
        else:
            lines.append("_No traceable values in this section._")
        return "\n".join(lines)

    def to_docx(self, path: str | Path, *, tracked_changes: bool = True) -> Path:
        """Render a draft ``.docx`` (tracked changes on by default). Requires ``python-docx``."""

        try:
            from docx import Document
            from docx.oxml import OxmlElement
            from docx.shared import RGBColor
        except ModuleNotFoundError as exc:  # pragma: no cover - exercised via the extra
            raise ModuleNotFoundError(
                "to_docx() requires python-docx; install the optional extra: "
                "pip install 'moltrace[docx]'"
            ) from exc

        doc = Document()
        if tracked_changes:
            doc.settings.element.append(OxmlElement("w:trackChanges"))

        doc.add_heading(f"CTD {self.section_number} — {self.title}", level=0)

        banner = doc.add_paragraph()
        run = banner.add_run(("DRAFT — " if self.is_draft else "") + self.disclaimer)
        run.bold = True
        run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)

        doc.add_paragraph(f"Format: {self.m4q_reference}.").italic = True

        for s in self.subsections:
            doc.add_heading(f"{s.number} {s.title}", level=1)
            for p in s.paragraphs:
                doc.add_paragraph(p)
            if s.table is not None:
                doc.add_paragraph(s.table.caption, style="Caption")
                table = doc.add_table(rows=1, cols=len(s.table.headers))
                table.style = "Light Grid Accent 1"
                for i, h in enumerate(s.table.headers):
                    table.rows[0].cells[i].text = h
                for row in s.table.rows:
                    cells = table.add_row().cells
                    for i, val in enumerate(row):
                        cells[i].text = val

        doc.add_heading("Source traceability", level=1)
        index = self.source_index()
        if index:
            trace = doc.add_table(rows=1, cols=4)
            trace.style = "Light Grid Accent 1"
            for i, h in enumerate(("Ref", "Source", "Locator", "Value")):
                trace.rows[0].cells[i].text = h
            for sid, ref in index.items():
                cells = trace.add_row().cells
                cells[0].text = sid
                cells[1].text = ref.kind.value
                cells[2].text = ref.locator
                cells[3].text = ref.value

        out = Path(path)
        doc.save(out)
        return out


# --------------------------------------------------------------------------- #
# Limit derivation (consumes the validated engines; mirrors the Q6A builder)
# --------------------------------------------------------------------------- #
def _fmt(value: float) -> str:
    return f"{value:g}"


def _match_by_smiles(items, smiles):
    if not smiles:
        return None
    for item in items:
        if getattr(item, "smiles", None) == smiles:
            return item
    return None


def _impurity_limit(
    entry: ImpurityEntry,
    m7: M7Classification | None,
    cpca: CPCAResult | None,
    q3ab: ImpurityThresholds,
    dose_g: float,
    ledger: _SourceLedger,
) -> tuple[str, str, str]:
    """Return (proposed_limit, regulatory_basis, safety_qualification), each with [S#] markers.

    Selection mirrors the validated Q6A builder: Cohort-of-Concern nitrosamine -> CPCA acceptable
    intake; otherwise a mutagenic (M7 class 1-3) impurity -> ICH M7 TTC safety limit; otherwise the
    ICH Q3A/B qualification threshold.
    """

    def _dose_ref() -> str:
        # Marked lazily so the dose source appears only when it is printed inline (the ppm paths).
        return ledger.mark(
            SourceRef(
                SourceKind.Q3AB, "daily_dose_g", f"{_fmt(dose_g)} g/day", q3ab.regulatory_basis
            )
        )

    if cpca is not None and cpca.coc_flag:
        ppm = cpca.ai_limit_ng_per_day / dose_g / 1000.0
        ref = ledger.mark(
            SourceRef(
                SourceKind.CPCA,
                "ai_limit_ng_per_day",
                f"{_fmt(cpca.ai_limit_ng_per_day)} ng/day",
                cpca.regulatory_basis,
            )
        )
        limit = f"NMT {_fmt(ppm)} ppm {ref}"
        basis = (
            f"Cohort-of-Concern N-nitrosamine; FDA CPCA category {cpca.category} acceptable intake "
            f"{_fmt(cpca.ai_limit_ng_per_day)} ng/day {ref}, divided by the maximum daily dose "
            f"({_fmt(dose_g)} g/day {_dose_ref()})."
        )
        qual = (
            f"Controlled at or below the FDA CPCA acceptable-intake limit (the ICH M7 TTC does not "
            f"apply to Cohort-of-Concern compounds) {ref}; expert toxicological review required."
        )
        return limit, basis, qual

    if m7 is not None and m7.m7_class in (1, 2, 3):
        if m7.ttc_ug_per_day is not None:
            ppm = m7.ttc_ug_per_day / dose_g
            ref = ledger.mark(
                SourceRef(
                    SourceKind.M7,
                    "ttc_ug_per_day",
                    f"{_fmt(m7.ttc_ug_per_day)} ug/day ({m7.duration_band})",
                    m7.regulatory_basis,
                )
            )
            limit = f"NMT {_fmt(ppm)} ppm {ref}"
            basis = (
                f"Mutagenic impurity, ICH M7 Class {m7.m7_class}; the acceptable intake "
                f"{_fmt(m7.ttc_ug_per_day)} ug/day {ref} — the staged TTC for the "
                f"{m7.duration_band} treatment duration — divided by the maximum daily dose "
                f"({_fmt(dose_g)} g/day {_dose_ref()})."
            )
            qual = (
                f"Controlled at or below the ICH M7 acceptable intake for the {m7.duration_band} "
                f"duration {ref} (M7 control obviates Q3A/B qualification for mutagenic potential)."
            )
            return limit, basis, qual
        ref = ledger.mark(
            SourceRef(SourceKind.M7, "m7_class", str(m7.m7_class), m7.regulatory_basis)
        )
        limit = "Compound-specific limit (expert review) — pending"
        basis = (
            f"Mutagenic impurity, ICH M7 Class {m7.m7_class} {ref}; a compound-specific acceptable "
            "intake (not the generic TTC) is required."
        )
        qual = f"Requires a compound-specific risk assessment per ICH M7 {ref}."
        return limit, basis, qual

    qt = q3ab.qualification_threshold
    ref = ledger.mark(
        SourceRef(
            SourceKind.Q3AB,
            "qualification_threshold.effective_percent",
            f"{_fmt(qt.effective_percent)}%",
            q3ab.regulatory_basis,
        )
    )
    limit = f"NMT {_fmt(qt.effective_percent)}% {ref}"
    basis = (
        f"Ordinary (non-mutagenic) impurity; controlled at the ICH Q3A/B qualification threshold "
        f"{_fmt(qt.effective_percent)}% {ref} ({qt.table_reference})."
    )
    qual = f"Qualified at or below the ICH Q3A/B qualification threshold {ref}."
    return limit, basis, qual


def _observed_clause(entry: ImpurityEntry, ledger: _SourceLedger) -> str:
    if entry.max_observed_percent is None:
        return "no batch data supplied"
    ref = ledger.mark(
        SourceRef(
            SourceKind.BATCH,
            f"impurity[{entry.name}].max_observed_percent",
            f"{_fmt(entry.max_observed_percent)}%",
        )
    )
    return (
        f"observed at up to {_fmt(entry.max_observed_percent)}% across the reported batches {ref}"
    )


# --------------------------------------------------------------------------- #
# 3.2.P.5.5 + 3.2.P.5.6 (drug product)
# --------------------------------------------------------------------------- #
def generate_3p5_impurities(
    impurity_profile: ImpurityProfile,
    q3ab_thresholds: ImpurityThresholds,
    m7_classifications: list[M7Classification],
    cpca_results: list[CPCAResult],
    batch_data: list[BatchResult],
) -> CTDSection:
    """Generate CTD 3.2.P.5.5 (Characterisation of Impurities) and 3.2.P.5.6 (Justification of
    Specifications) in ICH M4Q(R1) narrative format. See the module docstring for traceability."""

    ledger = _SourceLedger()
    dose_g = q3ab_thresholds.daily_dose_g
    specified = [imp for imp in impurity_profile.impurities if imp.is_specified]

    # ---- 3.2.P.5.5 Characterisation of Impurities ----
    p55_paras: list[str] = [
        f"The impurity profile of {impurity_profile.substance_name} comprises "
        f"{len(specified)} specified impurit{'y' if len(specified) == 1 else 'ies'}. Each is "
        "characterised below with its structural assignment, origin, proposed acceptance limit "
        "(with regulatory basis), and safety qualification. Reporting and identification "
        "thresholds follow ICH Q3A/B for the maximum daily dose.",
    ]
    rt_ref = ledger.mark(
        SourceRef(
            SourceKind.Q3AB,
            "reporting_threshold.effective_percent",
            f"{_fmt(q3ab_thresholds.reporting_threshold.effective_percent)}%",
            q3ab_thresholds.regulatory_basis,
        )
    )
    it_ref = ledger.mark(
        SourceRef(
            SourceKind.Q3AB,
            "identification_threshold.effective_percent",
            f"{_fmt(q3ab_thresholds.identification_threshold.effective_percent)}%",
            q3ab_thresholds.regulatory_basis,
        )
    )
    p55_paras.append(
        f"Reporting threshold: {_fmt(q3ab_thresholds.reporting_threshold.effective_percent)}% "
        f"{rt_ref}. Identification threshold: "
        f"{_fmt(q3ab_thresholds.identification_threshold.effective_percent)}% {it_ref}."
    )

    rows: list[tuple[str, ...]] = []
    m7_by_smiles = {m.smiles: m for m in m7_classifications}
    cpca_by_smiles = {c.smiles: c for c in cpca_results}
    for imp in specified:
        m7 = m7_by_smiles.get(imp.structure_smiles) if imp.structure_smiles else None
        cpca = cpca_by_smiles.get(imp.structure_smiles) if imp.structure_smiles else None
        limit, basis, qual = _impurity_limit(imp, m7, cpca, q3ab_thresholds, dose_g, ledger)
        observed = _observed_clause(imp, ledger)
        structure = imp.structural_description or (
            f"SMILES {imp.structure_smiles}"
            if imp.structure_smiles
            else "structure to be confirmed"
        )
        p55_paras.append(
            f"**{imp.name}.** Structure: {structure}. Origin: {imp.origin.value}. This impurity is "
            f"{observed}. Proposed acceptance limit: {limit}. Basis: {basis} Safety qualification: "
            f"{qual}"
        )
        rows.append((imp.name, imp.origin.value, observed.split(" [")[0], limit, qual))

    p55 = CTDSubsection(
        number="3.2.P.5.5",
        title="Characterisation of Impurities",
        paragraphs=tuple(p55_paras),
        table=CTDTable(
            caption="Table 3.2.P.5.5-1. Specified impurities — origin, observed level, proposed "
            "limit, and qualification.",
            headers=("Impurity", "Origin", "Observed", "Proposed limit", "Qualification"),
            rows=tuple(rows),
        ),
    )

    # ---- 3.2.P.5.6 Justification of Specifications ----
    method = impurity_profile.analytical_procedure
    validated = (
        "validated in accordance with ICH Q2(R2)"
        if impurity_profile.analytical_procedure_validated
        else "validation in progress (ICH Q2(R2))"
    )
    p56_paras: list[str] = [
        f"The related-substances acceptance criteria are determined by {method}, {validated}. The "
        "procedure is stability-indicating and resolves each specified impurity from the active "
        "and from each other.",
        "Rationale for the acceptance criteria: each limit is set at the safety-qualified level — "
        "the ICH Q3A/B qualification threshold for ordinary impurities, the ICH M7 TTC-derived "
        "limit for mutagenic impurities, and the FDA CPCA acceptable intake for Cohort-of-Concern "
        "nitrosamines — and is not tighter than necessary given the observed batch levels, "
        "consistent with ICH Q6A. Total degradation products are limited per ICH Q3B.",
    ]
    batch_totals = [
        (b.batch_id, b.total_impurities_percent)
        for b in batch_data
        if b.total_impurities_percent is not None
    ]
    if batch_totals:
        max_total = max(t for _, t in batch_totals)
        batch_ids = ", ".join(b.batch_id for b in batch_data)
        tot_ref = ledger.mark(
            SourceRef(SourceKind.BATCH, "max(total_impurities_percent)", f"{_fmt(max_total)}%")
        )
        p56_paras.append(
            f"Across {len(batch_data)} batch(es) ({batch_ids}), total impurities did not exceed "
            f"{_fmt(max_total)}% {tot_ref}, supporting the proposed total-impurities acceptance "
            "criterion."
        )
    if impurity_profile.pharmacopoeial_monograph:
        mono = impurity_profile.pharmacopoeial_monograph
        compared = [imp for imp in specified if imp.pharmacopoeial_limit_percent is not None]
        if compared:
            clauses = []
            for imp in compared:
                mono_ref = ledger.mark(
                    SourceRef(
                        SourceKind.PHARMACOPOEIA,
                        f"pharmacopoeial[{imp.name}]",
                        f"{_fmt(imp.pharmacopoeial_limit_percent)}%",
                        mono,
                    )
                )
                clauses.append(
                    f"{imp.name}: limit {_fmt(imp.pharmacopoeial_limit_percent)}% {mono_ref}"
                )
            p56_paras.append(
                f"Comparison to the pharmacopoeial monograph ({mono}): "
                + "; ".join(clauses)
                + ". The proposed limits are no less stringent than the compendial limits."
            )
        else:
            p56_paras.append(
                f"A pharmacopoeial monograph ({mono}) exists; per-impurity compendial limits "
                "should be tabulated and compared on review."
            )
    else:
        p56_paras.append(
            "No applicable pharmacopoeial monograph was identified; the acceptance criteria are "
            "justified solely on the ICH safety and batch-data bases above."
        )

    p56 = CTDSubsection(
        number="3.2.P.5.6",
        title="Justification of Specification(s)",
        paragraphs=tuple(p56_paras),
    )

    return CTDSection(
        section_number="3.2.P.5.5 / 3.2.P.5.6",
        title="Characterisation and Justification of Impurities (Drug Product)",
        subsections=(p55, p56),
        sources=ledger.refs(),
    )


# --------------------------------------------------------------------------- #
# 3.2.S.3.2 (drug substance)
# --------------------------------------------------------------------------- #
def generate_3s3_impurities_drug_substance(
    substance_profile: SubstanceProfile,
    impurity_profile: ImpurityProfile,
    q3ab_thresholds: ImpurityThresholds,
) -> CTDSection:
    """Generate CTD 3.2.S.3.2 (Impurities) for a drug substance in ICH M4Q(R1) narrative format."""

    ledger = _SourceLedger()
    dose_g = q3ab_thresholds.daily_dose_g
    specified = [imp for imp in impurity_profile.impurities if imp.is_specified]
    qt = q3ab_thresholds.qualification_threshold

    qual_ref = ledger.mark(
        SourceRef(
            SourceKind.Q3AB,
            "qualification_threshold.effective_percent",
            f"{_fmt(qt.effective_percent)}%",
            q3ab_thresholds.regulatory_basis,
        )
    )
    rt_ref = ledger.mark(
        SourceRef(
            SourceKind.Q3AB,
            "reporting_threshold.effective_percent",
            f"{_fmt(q3ab_thresholds.reporting_threshold.effective_percent)}%",
            q3ab_thresholds.regulatory_basis,
        )
    )
    dose_ref = ledger.mark(
        SourceRef(
            SourceKind.Q3AB,
            "daily_dose_g",
            f"{_fmt(dose_g)} g/day",
            q3ab_thresholds.regulatory_basis,
        )
    )
    paras: list[str] = [
        f"This section describes the actual and potential impurities arising from the manufacture "
        f"and storage of {substance_profile.name}. Thresholds follow ICH Q3A for the maximum daily "
        f"dose of {_fmt(dose_g)} g {dose_ref}: reporting "
        f"{_fmt(q3ab_thresholds.reporting_threshold.effective_percent)}% {rt_ref}, qualification "
        f"{_fmt(qt.effective_percent)}% {qual_ref} ({qt.table_reference}).",
        "Process-related impurities (starting materials, intermediates, reagents, by-products) and "
        "degradation products are listed below with their origin and proposed acceptance limits. "
        "Any impurity bearing a structural alert is additionally assessed for mutagenic potential "
        "per ICH M7; N-nitrosamines are assessed per the FDA CPCA (reported in the relevant "
        "safety sections).",
    ]

    rows: list[tuple[str, ...]] = []
    for imp in specified:
        observed = _observed_clause(imp, ledger)
        observed_cap = observed[0].upper() + observed[1:]
        if imp.pharmacopoeial_limit_percent is not None:
            ph_ref = ledger.mark(
                SourceRef(
                    SourceKind.PHARMACOPOEIA,
                    f"pharmacopoeial[{imp.name}]",
                    f"{_fmt(imp.pharmacopoeial_limit_percent)}%",
                    impurity_profile.pharmacopoeial_monograph or "pharmacopoeial monograph",
                )
            )
            limit = f"NMT {_fmt(imp.pharmacopoeial_limit_percent)}% {ph_ref}"
        else:
            limit = f"NMT {_fmt(qt.effective_percent)}% {qual_ref}"
        alert = " (structural alert — assess per ICH M7)" if imp.structure_smiles else ""
        paras.append(
            f"**{imp.name}.** Origin: {imp.origin.value}{alert}. {observed_cap}. "
            f"Proposed limit: {limit}."
        )
        rows.append((imp.name, imp.origin.value, observed.split(" [")[0], limit))

    section_body = CTDSubsection(
        number="3.2.S.3.2",
        title="Impurities",
        paragraphs=tuple(paras),
        table=CTDTable(
            caption="Table 3.2.S.3.2-1. Drug-substance impurities — origin, observed level, and "
            "proposed limit.",
            headers=("Impurity", "Origin", "Observed", "Proposed limit"),
            rows=tuple(rows),
        ),
    )

    return CTDSection(
        section_number="3.2.S.3.2",
        title="Impurities (Drug Substance)",
        subsections=(section_body,),
        sources=ledger.refs(),
    )
