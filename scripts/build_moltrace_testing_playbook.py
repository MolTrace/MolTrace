from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path("MolTrace_Testing_AI_ML_Playbook.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B677A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "E9F5EE"
PALE_GOLD = "FFF4D6"
PALE_RED = "FCE8E6"
BORDER = "C9D3DF"
WHITE = "FFFFFF"


def rgb(hex_color: str) -> RGBColor:
    hex_color = hex_color.lstrip("#")
    return RGBColor(
        int(hex_color[0:2], 16),
        int(hex_color[2:4], 16),
        int(hex_color[4:6], 16),
    )


def set_run_font(run, *, size: float | None = None, color: str | None = None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = BORDER, size: str = "6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def mark_row_as_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_table_width(table, widths):
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(cell)


def style_table(table, widths, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, widths)
    mark_row_as_header(table.rows[0])
    if header:
        for cell in table.rows[0].cells:
            shade_cell(cell, LIGHT_BLUE)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_run_font(r, size=9.5, color=INK, bold=True)
    for row in table.rows[1 if header else 0 :]:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    set_run_font(r, size=9.2, color="111111")


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=9, color=MUTED, italic=True)
    return p


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    style_table(table, [6.35], header=False)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=INK, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=9.8, color="222222")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color="111111")
    return p


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color="111111")
    return p


def add_small_note(doc: Document, text: str):
    p = doc.add_paragraph(style="Source Note")
    run = p.add_run(text)
    set_run_font(run, size=8.7, color=MUTED)
    return p


def add_code_block(doc: Document, text: str):
    table = doc.add_table(rows=1, cols=1)
    style_table(table, [6.35], header=False)
    cell = table.cell(0, 0)
    shade_cell(cell, "F7F9FB")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    run.font.size = Pt(8.5)
    run.font.color.rgb = rgb("263238")


def add_scheme_row(table, row_idx: int, labels: list[str], fills: list[str] | None = None):
    cells = table.rows[row_idx].cells
    fills = fills or [WHITE] * len(labels)
    for idx, label in enumerate(labels):
        cells[idx].text = ""
        p = cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        set_run_font(run, size=9.4, color=INK, bold=idx % 2 == 0)
        shade_cell(cells[idx], fills[idx])


def add_scheme_table(doc: Document, title: str, rows: list[list[str]], fills: list[list[str]], widths: list[float]):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    style_table(table, widths, header=False)
    for idx, row in enumerate(rows):
        add_scheme_row(table, idx, row, fills[idx])
    add_caption(doc, "Editable scheme table for " + title.split(".", 1)[-1].strip())
    return table


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Title", 24, INK, 0, 6),
        ("Subtitle", 13, MUTED, 0, 12),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        if style_name in {"Title", "Heading 1", "Heading 2", "Heading 3"}:
            style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        styles[list_style].font.name = "Calibri"
        styles[list_style].font.size = Pt(10.5)
        styles[list_style].paragraph_format.space_after = Pt(4)
        styles[list_style].paragraph_format.line_spacing = 1.25

    if "Caption" in styles:
        styles["Caption"].font.name = "Calibri"
        styles["Caption"].font.size = Pt(9)
        styles["Caption"].font.color.rgb = rgb(MUTED)
        styles["Caption"].paragraph_format.space_after = Pt(6)
    source_style = styles.add_style("Source Note", 1)
    source_style.font.name = "Calibri"
    source_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    source_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    source_style.font.size = Pt(8.7)
    source_style.font.color.rgb = rgb(MUTED)
    source_style.paragraph_format.space_before = Pt(2)
    source_style.paragraph_format.space_after = Pt(4)
    source_style.paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    header.text = ""
    header_run = header.add_run("MolTrace Testing and AI/ML Playbook | May 2026")
    set_run_font(header_run, size=8.5, color=MUTED)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer.paragraphs[0]
    footer.text = ""
    f_run = footer.add_run("Prepared for MolTrace product validation and native AI rollout")
    set_run_font(f_run, size=8.5, color=MUTED)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc


def add_cover(doc: Document):
    doc.add_paragraph().paragraph_format.space_after = Pt(18)
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    r = kicker.add_run("MOLTRACE PRODUCT VALIDATION GUIDE")
    set_run_font(r, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph(style="Title")
    title.paragraph_format.space_after = Pt(2)
    title.add_run("Testing, Troubleshooting, and Native AI/ML Rollout Playbook")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(
        "A practical guide for validating SpectraCheck first, then Regulatory Hub, then ReactionIQ, while turning reviewed evidence into governed AI training data."
    )

    rows = [
        ("Prepared", "May 18, 2026"),
        ("Product sequence", "SpectraCheck -> Regulatory Hub -> ReactionIQ"),
        ("Verified from", "MolTraceDocs, backend contracts, API surface, and current deployed OnRender frontend"),
        ("Primary use", "Pilot testing, batch testing, ML readiness, and operator troubleshooting"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    style_table(table, [1.55, 4.8], header=False)
    for idx, (label, value) in enumerate(rows):
        shade_cell(table.cell(idx, 0), LIGHT_BLUE)
        p0 = table.cell(idx, 0).paragraphs[0]
        p0.text = ""
        rr = p0.add_run(label)
        set_run_font(rr, size=9.5, color=INK, bold=True)
        p1 = table.cell(idx, 1).paragraphs[0]
        p1.text = ""
        rr1 = p1.add_run(value)
        set_run_font(rr1, size=9.5, color="111111")

    add_callout(
        doc,
        "Core recommendation",
        "Do not start by throwing millions of compounds at the system. Start with a locked golden SpectraCheck pilot, turn reviewed outputs into dataset candidates, then scale through chunked batch jobs with frozen holdout sets and model governance gates.",
        PALE_GOLD,
    )

    doc.add_page_break()


def add_contents(doc: Document):
    doc.add_heading("Contents", level=1)
    items = [
        "1. Product-readiness stance",
        "2. Step-by-step testing plan using available data",
        "3. Troubleshooting and iteration guide",
        "4. Native AI/ML operating model",
        "5. Batch and million-compound testing",
        "6. Recommended next steps",
        "Appendix A. Public data/resources",
        "Appendix B. Endpoint and artifact checklist",
        "Appendix C. Source list",
    ]
    for item in items:
        add_bullet(doc, item)
    doc.add_page_break()


def add_product_readiness(doc: Document):
    doc.add_heading("1. Product-Readiness Stance", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "MolTrace is pointed in the right direction: the local documentation and backend contract show a governance-first scientific platform rather than a single black-box chemistry app. The strongest current foundation is the evidence architecture: raw record preservation, SHA-256 provenance, review states, report packages, and model-governance records."
    )
    p = doc.add_paragraph()
    p.add_run(
        "The important caveat is that the production ML layer is not yet the same thing as a production-scale model-training stack. The current model factory and AI service surfaces are valuable because they define the gates: dataset versions, training runs, evaluation runs, model cards, deployment candidates, monitoring, feedback, and active learning. Large graph/transformer/external training still needs trainer integration, artifact storage, distributed execution, and frozen evaluation data before it can honestly be called production training."
    )

    add_scheme_table(
        doc,
        "Scheme 1. Recommended product validation sequence",
        [
            ["SpectraCheck", "->", "Regulatory Hub", "->", "ReactionIQ"],
            ["Identity evidence and spectra", "feeds", "rules, dossiers, readiness", "feeds", "reaction objectives and constraints"],
            ["Golden pilot and batch evidence", "gate", "compliance interpretation", "gate", "optimization loop"],
        ],
        [
            [LIGHT_BLUE, WHITE, LIGHT_BLUE, WHITE, LIGHT_BLUE],
            [WHITE, LIGHT_GRAY, WHITE, LIGHT_GRAY, WHITE],
            [PALE_GREEN, WHITE, PALE_GREEN, WHITE, PALE_GREEN],
        ],
        [1.35, 0.35, 1.55, 0.35, 1.55],
    )

    doc.add_heading("What is already testable", level=2)
    for item in [
        "1H NMR validation/analyze with SMILES, solvent-aware scoring, 13C beta validation, DEPT/APT-style evidence, processed spectrum previews, raw 1D FID beta vaulting/processing, and guarded 2D NMR support.",
        "Candidate comparison, spectral similarity, candidate-specific predicted NMR matching, HRMS/MS/MS, LC-MS import, feature grouping, and unified candidate confidence pathways.",
        "SpectraCheck sessions, evidence objects, review flows, reports, files, jobs, artifacts, reviewers, comments, tasks, approvals, and audit records.",
        "Regulatory sources, surveillance, rule updates, dossiers, evidence links, batch assessments, impurity/residual-solvent/nitrosamine/qNMR/method-validation panels, CTD Module 3 bundles, and AI-governance records.",
        "ReactionIQ projects, design spaces, objective/cost/safety profiles, Bayesian optimization, advisor runs, recommendations, execution batches, analytical results, outcome confirmation, and SpectraCheck evidence linking.",
        "ML/AI dashboards for tasks, feature pipelines, training runs, evaluation runs, artifacts, model cards, calibration/OOD/error analysis, deployment candidates, AI services, predictions, feedback, active learning, shadow/canary evaluations, and monitoring.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Immediate constraint to respect", level=2)
    add_callout(
        doc,
        "Batch limit today",
        "The legacy batch analysis model currently caps a submitted job at 100 items. That is enough for smoke tests and early pilot batches, but not for millions of compounds. Million-scale testing should use manifest-driven chunking, queues/workers, object storage, resumability, idempotency, and frozen evaluation splits.",
        PALE_RED,
    )


def add_testing_plan(doc: Document):
    doc.add_heading("2. Step-by-Step Testing Plan Using Available Data", level=1)
    add_callout(
        doc,
        "Phase order",
        "Start with SpectraCheck because it creates the evidence backbone. Move to Regulatory Hub after evidence objects are reliable. Move to ReactionIQ only after analytical evidence and regulatory constraints can be linked to experiments.",
        PALE_GREEN,
    )

    doc.add_heading("Phase 0 - Prepare the test environment", level=2)
    steps = [
        "Freeze one test environment with explicit backend URL, frontend URL, database, object/vault storage, Redis/queue configuration, admin users, and feature flags. Record the exact commit or deployment revision.",
        "Create a test data inventory with sample ID, compound name, SMILES/InChIKey, 1H text, 13C text, MS/LC-MS file or peak table, solvent, instrument/vendor, truth label, source, expected warnings, and reviewer signoff status.",
        "Separate data into three buckets: tiny smoke set, golden pilot set, and holdout set. The holdout set must not be used for tuning thresholds or prompts.",
        "Turn on logging around request IDs, input hashes, job IDs, evidence IDs, reviewer actions, model artifact IDs, and export package hashes.",
        "Define acceptance thresholds before looking at results: parse success, evidence agreement, rank-1 accuracy, contradiction detection, report completeness, runtime, and review turnaround.",
    ]
    for step in steps:
        add_number(doc, step)

    doc.add_heading("Phase 1 - SpectraCheck individual tests", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Begin with 20 individual compounds. Include easy positives, close analogs, invalid inputs, ambiguous spectra, solvents, exchangeable protons, missing 13C, and at least a few MS/LC-MS-supported examples. Run every case manually through the UI and API so the team can see how evidence, warnings, and review states behave."
    )
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    for idx, text in enumerate(["Test", "Input", "Expected signal", "Evidence to inspect", "Pass condition"]):
        hdr[idx].text = text
    rows = [
        ("1H only", "SMILES + 1H text", "Parsed peaks and agreement score", "matched/unmatched peaks, solvent warnings", "Score matches expectation and review is required"),
        ("1H + 13C", "SMILES + 1H + 13C", "Cross-layer evidence", "carbon count, region checks, contradictions", "No hidden contradiction; limitations visible"),
        ("Candidate ranking", "3-10 candidates + spectra", "Best-supported ranking", "ranked candidates, evidence_label", "Known structure ranks correctly or ambiguity is explained"),
        ("Processed spectra", "CSV/TSV/JDX/DX", "Preview and analyze", "axis/reference matching, unsupported formats", "Good files pass; bad files fail clearly"),
        ("Raw FID beta", "Bruker/Varian archive", "Vault + derivative run", "SHA-256, recipe, processing metadata", "Raw file immutable; derived package auditable"),
        ("2D beta", "HSQC/HMQC table", "Direct C-H context", "2D similarity, unsupported COSY/HMBC prediction limits", "2D support improves or flags ambiguity"),
        ("MS/LC-MS", "HRMS/MS/MS/LC-MS tables", "Orthogonal evidence", "adduct, isotope, blank, feature-family consensus", "MS evidence changes confidence without hiding warnings"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
    style_table(table, [0.9, 1.2, 1.35, 1.55, 1.25])

    doc.add_heading("Phase 2 - SpectraCheck batch pilot", level=2)
    for step in [
        "Run the same golden cases through batch mode in chunks of 25-100 so individual and batch outputs can be compared one-to-one.",
        "Export CSV/JSON job results and compare: sample_id, score, evidence_label, warning count, contradiction count, runtime, and generated evidence/report IDs.",
        "Measure drift between manual and batch paths. Any difference must be explained by input normalization, feature flag state, or asynchronous job metadata.",
        "Add negative controls: invalid SMILES, empty NMR text, malformed peak text, unsupported files, stale auth, duplicate sample IDs, and corrupted FID archives.",
        "Promote only reviewed, traceable cases into Knowledge records or dataset candidates.",
    ]:
        add_number(doc, step)

    doc.add_heading("Phase 3 - Regulatory Hub tests", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Regulatory Hub testing should begin after SpectraCheck evidence can be created, reviewed, exported, and linked. The first goal is not broad global regulatory coverage; it is traceable transformation from source rule to dossier requirement to evidence link to readiness/risk output."
    )
    table = doc.add_table(rows=1, cols=4)
    for idx, text in enumerate(["Area", "Test action", "Expected artifact", "Review question"]):
        table.rows[0].cells[idx].text = text
    rows = [
        ("Source library", "Upload/search official sources and versions", "source, citations, version timeline", "Can every rule point to a source?"),
        ("Dossiers", "Create dossier and requirements", "requirement records and evidence links", "Is evidence linked by ID, not pasted text?"),
        ("Batch assessment", "Run controlled compound batch", "risk/readiness output", "Are thresholds explainable and versioned?"),
        ("AI governance", "Create AI governance record", "model/use/validation context", "Would FDA/EMA reviewer see lineage?"),
        ("Submission package", "Generate CTD/readiness package", "exported package", "Can source, evidence, review, and hash chain be inspected?"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
    style_table(table, [1.15, 1.85, 1.55, 1.65])

    doc.add_heading("Phase 4 - ReactionIQ tests", level=2)
    for step in [
        "Create a small reaction project with a bounded design space, objective profile, cost profile, safety profile, and explicit regulatory constraints.",
        "Seed baseline experiments and analytical results, then link SpectraCheck sessions/evidence to each experiment.",
        "Run a Bayesian optimization recommendation and compare it against a human chemist's next-experiment choice. Do not auto-approve.",
        "Execute one recommendation, upload analytical output, extract/confirm the reaction outcome, and update the optimization cycle.",
        "Track whether ReactionIQ uses constraints and evidence correctly: yield, purity, cost, safety, solvent/reagent constraints, impurity flags, and confidence/uncertainty.",
    ]:
        add_number(doc, step)


def add_troubleshooting(doc: Document):
    doc.add_heading("3. Troubleshooting and Iteration Guide", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Troubleshooting should follow the evidence chain, not the screen where the symptom appears. Most failures fall into one of seven layers: input integrity, parsing, chemistry expectation, scoring, persistence, review/export, or ML/AI routing."
    )

    table = doc.add_table(rows=1, cols=4)
    for idx, text in enumerate(["Symptom", "Likely layer", "First checks", "Iteration move"]):
        table.rows[0].cells[idx].text = text
    rows = [
        ("Correct compound ranks poorly", "scoring or parser", "peak parsing, solvent, expected shifts, missing 13C/MS", "add curated case; tune thresholds only against training split"),
        ("Everything looks high confidence", "calibration", "score distribution, negative controls, close analogs", "add contradiction-heavy controls and calibration plots"),
        ("Batch differs from individual", "normalization/job path", "payload serialization, feature flags, async metadata", "create one-to-one regression fixture"),
        ("FID report is suspect", "raw vault/processing", "SHA-256 verification, recipe, phase/baseline metadata", "reprocess as new derivative, never overwrite raw archive"),
        ("Regulatory output cites weak source", "source/retrieval", "source version, citation support, jurisdiction", "require source review before dossier use"),
        ("ReactionIQ recommends unsafe design", "constraints/objectives", "safety profile, regulatory constraints, objective weights", "block auto-approval; create constraint regression"),
        ("AI prediction has no artifact lineage", "AI routing", "approved artifact, service config, model card, deployment candidate", "fail closed until approved model exists"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
    style_table(table, [1.45, 1.15, 1.95, 1.65])

    doc.add_heading("Iteration cadence", level=2)
    for item in [
        "Daily during pilot: triage failed cases into parser, chemistry/scoring, UI/API, data quality, governance, or model issue.",
        "Weekly: freeze a new golden set version, rerun all current tests, publish pass/fail deltas, and lock unresolved failures with owners.",
        "Before each threshold change: prove the change improves training/dev data and does not degrade the frozen holdout set.",
        "Before each AI deployment: require model card, evaluation run, calibration/OOD analysis, error analysis, deployment-candidate approval, and rollback criteria.",
    ]:
        add_bullet(doc, item)

    add_callout(
        doc,
        "Rule of thumb",
        "If a result cannot explain its inputs, method context, limitations, warnings, review state, and artifact lineage, it is not ready for regulatory or training use.",
        PALE_GOLD,
    )


def add_ml_ai(doc: Document):
    doc.add_heading("4. Native AI/ML Operating Model", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "MolTrace becomes a native AI system when AI is embedded in the governed evidence lifecycle: predictions are made through approved model artifacts, uncertainty is visible, human review is captured, low-confidence cases create active-learning candidates, and reviewed evidence improves future datasets. The AI layer should not bypass SpectraCheck, Regulatory Hub, or ReactionIQ; it should learn from them."
    )

    add_scheme_table(
        doc,
        "Scheme 2. Native AI flywheel",
        [
            ["Reviewed evidence", "->", "Knowledge records", "->", "Dataset candidates"],
            ["Dataset versions", "->", "Training runs", "->", "Model artifacts"],
            ["Evaluation + model card", "->", "Deployment candidate", "->", "AI service predictions"],
            ["Feedback + active learning", "->", "New reviewed evidence", "->", "Next dataset version"],
        ],
        [
            [PALE_GREEN, WHITE, LIGHT_BLUE, WHITE, LIGHT_BLUE],
            [LIGHT_BLUE, WHITE, PALE_GREEN, WHITE, LIGHT_BLUE],
            [PALE_GOLD, WHITE, LIGHT_BLUE, WHITE, PALE_GREEN],
            [LIGHT_GRAY, WHITE, PALE_GREEN, WHITE, LIGHT_BLUE],
        ],
        [1.35, 0.32, 1.35, 0.32, 1.55],
    )

    doc.add_heading("How the ML part starts working", level=2)
    for step in [
        "Create reviewed SpectraCheck outputs. Only use evidence with stable sample IDs, source hashes, method context, warnings, limitations, and human review status.",
        "Promote trusted outputs into Knowledge records and dataset candidates. Record data source, license, split assignment, label provenance, and reviewer.",
        "Approve a dataset version. This is the first moment the system has a stable training input.",
        "Run an ML training run from the ML Model Factory using a built-in task key, dataset_version_id, feature_pipeline_id if applicable, model_family, parameters, and notes.",
        "Inspect whether the run produced model_artifact_id. For baseline/sklearn families, the current store can compute training metrics and create an artifact. For graph neural network, transformer, or external families, the current code marks the run as requiring review or integration rather than silently pretending full training occurred.",
        "Create evaluation, calibration, OOD, and error-analysis records against frozen holdout data.",
        "Create and approve a deployment candidate only after evaluation succeeds. Then configure AI Services to use the approved artifact.",
        "Route predictions through AI Services. Predictions should carry service_key, artifact/model context, confidence, uncertainty/OOD status, warnings, explanation links, monitoring events, and feedback/review hooks.",
    ]:
        add_number(doc, step)

    doc.add_heading("How you know training is really happening", level=2)
    table = doc.add_table(rows=1, cols=3)
    for idx, text in enumerate(["Signal", "Where to verify", "Meaning"]):
        table.rows[0].cells[idx].text = text
    rows = [
        ("dataset_version_id exists and is approved", "/knowledge and /ml dataset flow", "Training input is versioned and reviewed"),
        ("training run status is succeeded", "/ml/training-runs", "Model Factory completed the run path"),
        ("model_artifact_id is present", "/ml/model-artifacts/{id}", "An artifact record exists for this run"),
        ("evaluation run is succeeded", "/ml/evaluation-runs", "The artifact was tested against evaluation data"),
        ("model card exists", "/ml/model-cards", "Intended use, limitations, data, and metrics are documented"),
        ("calibration/OOD/error analysis exist", "/ml/calibration-assessments, /ml/ood-assessments, /ml/error-analysis", "Risk and uncertainty were assessed"),
        ("deployment candidate approved", "/ml/deployment-candidates/{id}/approve", "Human governance allowed deployment"),
        ("AI prediction cites artifact/service", "/ai/predictions", "Live inference is using governed model context"),
        ("feedback and active-learning records accumulate", "/ai/predictions/{id}/feedback and /ai/active-learning/candidates", "The system is learning from real use"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
    style_table(table, [1.75, 2.05, 2.4])

    doc.add_heading("Native AI guardrails", level=2)
    for item in [
        "Never train on unreviewed production predictions as if they were truth labels.",
        "Keep train/dev/test/holdout splits compound-disjoint when measuring structure/ranking generalization.",
        "Preserve out-of-distribution and low-confidence cases for active learning, not quiet deletion.",
        "For regulatory claims, pair model outputs with source citations, human review, intended-use boundaries, and audit logs.",
        "Make every model deployment reversible with shadow/canary evaluation before full promotion.",
    ]:
        add_bullet(doc, item)


def add_batch(doc: Document):
    doc.add_heading("5. Batch and Million-Compound Testing", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Batch testing must support both individual forensic review and high-throughput model improvement. Treat the current 100-item job cap as a useful pilot constraint, then build a larger batch-testing system around manifests, chunking, workers, object storage, and resumable evaluation."
    )

    add_scheme_table(
        doc,
        "Scheme 3. Scalable batch-testing architecture",
        [
            ["Data manifest", "->", "Chunk planner", "->", "Job queue"],
            ["Workers", "->", "Evidence store", "->", "Review queue"],
            ["Exports + metrics", "->", "Dataset candidates", "->", "Training/evaluation"],
        ],
        [
            [LIGHT_BLUE, WHITE, LIGHT_BLUE, WHITE, LIGHT_BLUE],
            [PALE_GREEN, WHITE, LIGHT_BLUE, WHITE, PALE_GOLD],
            [LIGHT_GRAY, WHITE, PALE_GREEN, WHITE, LIGHT_BLUE],
        ],
        [1.3, 0.32, 1.35, 0.32, 1.55],
    )

    doc.add_heading("Pilot batch design", level=2)
    for step in [
        "Run 100 items as five chunks of 20, then as one 100-item job. Confirm identical outputs.",
        "Run 500 items as five 100-item jobs. Measure failures, runtime, queue behavior, memory, duplicate handling, and export reliability.",
        "Create a dashboard that compares individual-vs-batch results by sample_id and evidence_id.",
        "Add per-job immutable metadata: manifest hash, code/deployment version, feature flags, dataset version, chunk number, retry count, and owner.",
    ]:
        add_number(doc, step)

    doc.add_heading("Million-scale requirements", level=2)
    table = doc.add_table(rows=1, cols=3)
    for idx, text in enumerate(["Requirement", "Why it matters", "Practical implementation"]):
        table.rows[0].cells[idx].text = text
    rows = [
        ("Manifest-first ingestion", "Millions of rows cannot be managed as ad hoc UI submissions", "Parquet/CSV manifest with stable IDs, paths, hashes, truth labels, license/source"),
        ("Chunking and idempotency", "Jobs will fail and resume", "chunk_id + item_id unique key; reruns replace only failed derivatives"),
        ("Object/artifact storage", "Spectra and results are too large for only relational rows", "store raw and derived artifacts separately; DB stores references and hashes"),
        ("Queue/workers", "Long jobs need parallelism and backpressure", "Redis/RQ/Celery/Temporal-style workers, rate limits, retries"),
        ("Frozen holdout", "Model improvement needs honest measurement", "never tune on holdout; report all changes against it"),
        ("Active learning loop", "Not every compound needs equal review", "prioritize low confidence, OOD, disagreement, and high-impact chemical classes"),
        ("Audit exports", "Regulated users need traceability", "export job manifest, inputs, outputs, warnings, review, metrics, and hashes"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
    style_table(table, [1.55, 2.15, 2.5])

    doc.add_heading("Metrics to track", level=2)
    for item in [
        "Parse success rate by modality: 1H, 13C, MS/MS, LC-MS, raw FID, 2D.",
        "Rank-1, top-3, and mean reciprocal rank for candidate ranking.",
        "Contradiction detection sensitivity on known wrong-candidate controls.",
        "Calibration: confidence bins versus observed correctness.",
        "OOD rate by compound class, instrument source, solvent, and modality.",
        "Batch throughput, retry rate, p50/p95 runtime, memory, and queue latency.",
        "Reviewer agreement, correction rate, feedback volume, and active-learning conversion.",
    ]:
        add_bullet(doc, item)


def add_next_steps(doc: Document):
    doc.add_heading("6. Recommended Next Steps", level=1)
    table = doc.add_table(rows=1, cols=4)
    for idx, text in enumerate(["When", "Priority", "Action", "Definition of done"]):
        table.rows[0].cells[idx].text = text
    rows = [
        ("Week 1", "High", "Lock the golden SpectraCheck pilot set and exact acceptance criteria", "20-50 reviewed cases with source hashes, truth labels, expected outputs, and holdout split"),
        ("Week 1", "High", "Run individual UI/API tests and capture evidence/report artifacts", "Every case has pass/fail, screenshots/export, evidence ID, reviewer note"),
        ("Week 2", "High", "Run batch parity tests at 25, 100, and 500 items", "Individual and batch paths agree; exports are stable"),
        ("Week 2", "High", "Create troubleshooting ledger and regression fixtures for every failure", "Each failure has owner, layer, fix, and rerun result"),
        ("Week 3", "Medium", "Promote reviewed evidence to Knowledge and dataset candidates", "Dataset candidate records include label provenance and split assignment"),
        ("Week 3", "Medium", "Run first governed baseline training/evaluation cycle", "training_run, artifact, evaluation, model card, calibration/OOD notes exist"),
        ("Week 4", "Medium", "Pilot Regulatory Hub with linked evidence and official sources", "Dossier readiness/risk outputs cite source records and evidence IDs"),
        ("Week 4+", "Medium", "Pilot ReactionIQ on one constrained reaction program", "BO/advisor recommendation links constraints, experiments, analytical outcome, and review decision"),
        ("Month 2", "High", "Build million-scale batch architecture", "manifest/chunk/job/result tables, object storage, worker scaling, frozen benchmark runner"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
    style_table(table, [0.75, 0.75, 2.4, 2.3])

    add_callout(
        doc,
        "Best next move",
        "The best immediate phase is not broad AI training. It is a disciplined SpectraCheck golden pilot plus batch parity testing. That creates the trustworthy labels and evidence lineage the native AI layer needs.",
        PALE_GREEN,
    )


def add_appendices(doc: Document):
    doc.add_page_break()
    doc.add_heading("Appendix A. Public Data and Resource Map", level=1)
    table = doc.add_table(rows=1, cols=4)
    for idx, text in enumerate(["Resource", "Best use", "Data type", "Caution"]):
        table.rows[0].cells[idx].text = text
    rows = [
        ("nmrXiv", "Open NMR records and FAIR-oriented metadata", "NMR spectra/metadata", "Confirm licenses, formats, and quality before training"),
        ("NMRShiftDB", "NMR shift reference and candidate checks", "1H/13C shift data", "Normalize structures and solvents carefully"),
        ("MassBank", "MS/MS reference spectra", "MS/MS spectra", "Instrument/adduct differences matter"),
        ("GNPS", "Molecular networking and MS/MS datasets", "MS/MS/LC-MS data", "Community data quality varies"),
        ("PubChem", "Compound identifiers, SMILES/InChIKey, properties", "compound metadata", "Not a truth source for spectra by itself"),
        ("Internal lab data", "Highest-value training and validation data", "spectra, methods, outcomes", "Needs consent, IP controls, review, and data lineage"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
    style_table(table, [1.25, 2.0, 1.35, 1.6])

    doc.add_heading("Appendix B. Endpoint and Artifact Checklist", level=1)
    for item in [
        "SpectraCheck: /prediction/nmr/match/evidence, /spectracheck/sessions, /spectracheck/sessions/{id}/evidence, /spectracheck/sessions/{id}/review, /spectracheck/sessions/{id}/reports, /jobs, /jobs/{id}/export.csv, /jobs/{id}/export.json.",
        "Raw FID: /raw-fid/upload, /raw-fid/{archive_id}, /raw-fid/{archive_id}/download, /raw-fid/{archive_id}/preview, /raw-fid/{archive_id}/process, /raw-fid/{archive_id}/export.",
        "Regulatory Hub: /regulatory/sources/upload, /regulatory/sources/search, /regulatory/dossiers, /regulatory/dossiers/{id}/evidence-links, /regulatory/dossiers/{id}/batch-assessment, /regulatory/dossiers/{id}/ai-governance-record, /regulatory/dossiers/{id}/ctd-module3-bundle.",
        "ReactionIQ: /reaction-projects, /reaction-projects/{id}/design-space, /reaction-projects/{id}/optimization/bo/run, /reaction-projects/{id}/advisor/run, /reaction-recommendations/{id}/approve, /reaction-execution-items/{id}/analytical-results, /reaction-experiments/{id}/link-spectracheck-session.",
        "ML: /ml/tasks, /ml/feature-pipelines, /ml/training-runs, /ml/evaluation-runs, /ml/model-artifacts, /ml/model-cards, /ml/calibration-assessments, /ml/ood-assessments, /ml/deployment-candidates, /ml/model-health.",
        "AI: /ai/services, /ai/predictions, /ai/predictions/{id}/feedback, /ai/predictions/{id}/review, /ai/routing/decide, /ai/active-learning/candidates, /ai/shadow-evaluations, /ai/canary-deployments, /ai/model-monitoring.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Appendix C. Source List", level=1)
    sources = [
        ("MolTrace Technical White Paper", "local: moltrace_docs/src/content/docs/moltrace-ecosystem-technical-white-paper.md"),
        ("MolTrace SpectraCheck backend contract", "local: moltrace_backend/docs/moltrace_spectracheck_backend_contract.md"),
        ("NMRCheck/MolTrace backend README", "local: moltrace_backend/README.md"),
        ("FDA draft guidance on AI in drug/biological regulatory decision-making", "https://www.fda.gov/regulatory-information/search-fda-guidance-documents/considerations-use-artificial-intelligence-support-regulatory-decision-making-drug-and-biological"),
        ("FDA Guiding Principles for Good AI Practice in Drug Development", "https://www.fda.gov/about-fda/artificial-intelligence-drug-development/guiding-principles-good-ai-practice-drug-development"),
        ("EMA reflection paper on AI in the medicinal product lifecycle", "https://www.ema.europa.eu/en/use-artificial-intelligence-ai-medicinal-product-lifecycle-scientific-guideline"),
        ("ICH Q2(R2) Validation of Analytical Procedures", "https://database.ich.org/sites/default/files/ICH_Q2%28R2%29_Guideline_2023_1130.pdf"),
        ("ICH Q14 Analytical Procedure Development", "https://database.ich.org/sites/default/files/ICH_Q14_Guideline_2023_1116.pdf"),
        ("nmrXiv documentation", "https://docs.nmrxiv.org/"),
        ("NMRShiftDB", "https://nmrshiftdb.nmr.uni-koeln.de/"),
        ("MassBank documentation", "https://massbank.github.io/MassBank-documentation/about.html"),
        ("GNPS documentation", "https://gnpsdocs.readthedocs.io/en/latest/"),
        ("PubChem PUG-REST", "https://pubchem.ncbi.nlm.nih.gov/docs/pug-rest"),
        ("Huang et al., routine NMR-based structure elucidation with artificial intelligence", "https://pubs.rsc.org/en/content/articlelanding/2021/sc/d1sc04105c"),
        ("Shields et al., Bayesian reaction optimization", "https://www.nature.com/articles/s41586-021-03213-y"),
        ("Chemical Society Reviews 2026 review on Bayesian optimization in chemical synthesis", "https://pubs.rsc.org/en/content/articlehtml/2026/cs/d5cs00962f"),
    ]
    for title, url in sources:
        p = doc.add_paragraph(style="Source Note")
        run = p.add_run(title + ": ")
        set_run_font(run, size=8.7, color="111111", bold=True)
        if url.startswith("http"):
            add_hyperlink(p, "Open source", url)
        else:
            run2 = p.add_run(url)
            set_run_font(run2, size=8.7, color=MUTED)


def build():
    doc = setup_document()
    add_cover(doc)
    add_contents(doc)
    add_product_readiness(doc)
    add_testing_plan(doc)
    add_troubleshooting(doc)
    add_ml_ai(doc)
    add_batch(doc)
    add_next_steps(doc)
    add_appendices(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH.resolve())


if __name__ == "__main__":
    build()
